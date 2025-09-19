from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from dataclasses import dataclass
from typing import List, Optional
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import shutil
from docx.shared import Inches
from datetime import datetime
from jinja2 import Template
import xml.etree.ElementTree as ET
from models import *
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocxTemplateGenerator:
 
    def __init__(self, template_path="Risk_Assessment_Form_Template.docx"):
        self.template_path = template_path
        
    def _center_align_columns(self, row, column_indices):
        """Center align text in specific columns"""
        for col_idx in column_indices:
            if col_idx < len(row.cells):
                for paragraph in row.cells[col_idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def generate_with_python_docx(self, assessment_id: str, form_data: dict = None, output_path: str = None):
        """Generate document using template with dynamic risk assessment table"""
        try:
            if output_path is None:
                current_dir = os.path.dirname(os.path.abspath(__file__))
                output_path = os.path.join(current_dir, f"risk_assessment_{assessment_id}.docx")
        
            # Load the template
            doc = self._load_template()
            if not doc:
                return None
            
            # Debug: Print the incoming form data
            print("=== DEBUG: Incoming form_data ===")
            print(f"form_data type: {type(form_data)}")
            print(f"form_data content: {form_data}")
            print("=== END DEBUG ===")

            # Check for team_data in form_data
            if form_data and 'form' in form_data and 'team_data' in form_data['form']:
                print(f"=== DEBUG: team_data found ===")
                print(f"team_data: {form_data['form']['team_data']}")
            else:
                print("=== DEBUG: team_data NOT found in form_data['form'] ===")
            
            # Use provided form_data or fallback to sample data
            if form_data is None or not form_data:
                print("No form_data provided")
                # Keep your original sample data as fallback
            else:
                print("Transforming form_data")

                print(f"DEBUG: 'division' in form_data: {'division' in form_data}")
                print(f"DEBUG: form_data.get('division'): {form_data.get('division')}")
                print(f"DEBUG: 'divisionName' in form_data: {'divisionName' in form_data}")
                print(f"DEBUG: form_data.get('divisionName'): {form_data.get('divisionName')}")

                # Transform the form data into the expected structure
                assessment_data = self._transform_form_data(form_data, assessment_id)
                
                # Debug: Print the transformed data
                print("=== DEBUG: Transformed assessment_data ===")
                print(f"assessment_data: {assessment_data}")
                print("=== END DEBUG ===")
            
            # Process the template
            self._process_template_with_risk_data(doc, assessment_data)
            
            self._replace_page_placeholders_with_fields(doc)

            # Save the document
            doc.save(output_path)
            print(f"Risk assessment document generated: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"Error generating document: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _format_possible_injuries(self, injury_text: str) -> str:
        """Format comma-separated injuries as a) b) c) list"""
        if not injury_text:
            return ''
        
        # Split by comma and clean up each injury
        injuries = [injury.strip() for injury in injury_text.split(',') if injury.strip()]
        
        # If only one injury, return as is
        if len(injuries) <= 1:
            return injury_text
        
        # Format as a) b) c) list
        formatted_injuries = []
        for i, injury in enumerate(injuries):
            letter = chr(ord('a') + i)  # Convert 0->a, 1->b, 2->c, etc.
            formatted_injuries.append(f"{letter}) {injury}")
        
        return '\n'.join(formatted_injuries)

    def _transform_form_data(self, form_data: dict, assessment_id: str) -> dict:
        """Transform form data into the expected assessment data structure"""
        try:
            print(f"=== TRANSFORM DEBUG: form_data keys: {list(form_data.keys()) if form_data else 'None'}")
            
            # Extract form info and activities data
            form_info = form_data.get('form', {})
            activities_data = form_data.get('activities_data', [])
            
            print(f"=== TRANSFORM DEBUGG: form_info: {form_info}")
            print(f"=== TRANSFORM DEBUG: activities_data count: {len(activities_data)}")
            

            def format_date(date_str):
                if not date_str or date_str == 'N/A':
                    return ''
                try:
                    # Parse the ISO format date
                    date_obj = datetime.fromisoformat(date_str.replace('T00:00:00', ''))
                    # Format as "14 Jul 2025"
                    return date_obj.strftime('%d %b %Y')
                except (ValueError, AttributeError):
                    return date_str  # Return original if parsing fails
                
            # Extract basic info from form data
            basic_info = {
                'assessment_id': assessment_id,
                'division': form_data.get('divisionName', 'N/A'),
                'department': form_info.get('department', 'N/A'),
                'location': form_info.get('location', 'N/A'),
                'supervisor': form_info.get('supervisor', 'N/A'),
                'date_created': format_date(form_info.get('last_review_date', 'N/A')),
                'next_review': format_date(form_info.get('next_review_date', 'N/A')),
                'title': form_info.get('title', f'Risk Assessment {assessment_id}')
            }

            # Add team_data if present
            if 'team_data' in form_info and isinstance(form_info['team_data'], dict):
                basic_info['team_data'] = form_info['team_data']
            
            print(f"=== TRANSFORM DEBUG: basic_info: {basic_info}")
            
            # Transform activities data into work_activities and processes
            work_activities = []
            processes = []
            
            # Group activities by process for the processes section
            process_groups = {}
            
            for activity in activities_data:
                # Add to work activities
                work_activity = {
                    'location': activity.get('location', ''),
                    'process': activity.get('process', ''),
                    'work_activity': activity.get('work_activity', ''),
                    'remarks': activity.get('remarks', '')
                }
                work_activities.append(work_activity)
                
                # Group by process for risk assessment
                process_name = activity.get('process', 'Unknown Process')
                if process_name not in process_groups:
                    process_groups[process_name] = {
                        'process_name': process_name,
                        'risks': []
                    }
                
                # Transform hazards into risks
                hazards = activity.get('hazards', [])
                for i, hazard in enumerate(hazards):
                    # Skip empty hazards
                    if not hazard.get('hazard') and not hazard.get('injury'):
                        continue
                        
                    risk = {
                        'ref': f"{len(process_groups)}.{len(process_groups[process_name]['risks']) + 1}",
                        'activity': activity.get('work_activity', ''),
                        'hazard': hazard.get('hazard', ''),
                        'possible_injury': self._format_possible_injuries(hazard.get('injury', '')),
                        'existing_controls': hazard.get('existing_controls', ''),
                        's1': str(hazard.get('severity', '')),
                        'l1': str(hazard.get('likelihood', '')),
                        'rpn1': str(hazard.get('rpn', '')),
                        'additional_controls': hazard.get('additional_controls', ''),
                        's2': str(hazard.get('newSeverity') or hazard.get('newseverity', '')),
                        'l2': str(hazard.get('newLikelihood') or hazard.get('newlikelihood', '')),
                        'rpn2': str(hazard.get('newRPN') or hazard.get('newrpn', '')),
                        'implementation_person': hazard.get('hazard_implementation_person', ''),
                        'due_date': format_date(hazard.get('hazard_due_date', '')),
                        'remarks': activity.get('remarks', '')
                    }
                    process_groups[process_name]['risks'].append(risk)
            
            # Convert process groups to list
            processes = list(process_groups.values())
            
            # Return transformed data structure
            transformed_data = {
                'basic_info': basic_info,
                'work_activities': work_activities,
                'processes': processes
            }
            
            print(f"=== TRANSFORM DEBUG: Final transformed data structure:")
            print(f"  - basic_info: {basic_info}")
            print(f"  - work_activities count: {len(work_activities)}")
            print(f"  - processes count: {len(processes)}")
            for process in processes:
                print(f"    - Process '{process['process_name']}': {len(process['risks'])} risks")
            
            return transformed_data
            
        except Exception as e:
            print(f"Error transforming form data: {e}")
            import traceback
            traceback.print_exc()
            # Return basic structure with assessment_id if transformation fails
            return {
                'basic_info': {
                    'assessment_id': assessment_id,
                    'division': 'Error',
                    'department': 'Error',
                    'location': 'Error',
                    'supervisor': 'Error',
                    'date_created': 'Error',
                    'title': f'Risk Assessment {assessment_id} - Error'
                },
                'work_activities': [],
                'processes': []
            }
    
    def _load_template(self):
        """Load the template document"""
        if not self.template_path:
            print("No template path provided")
            return None
        
        template_paths_to_try = [
            self.template_path,
            os.path.join(os.getcwd(), self.template_path),
            os.path.join('templates', self.template_path),
            os.path.join('app', 'templates', self.template_path),
            os.path.join(os.path.dirname(__file__), self.template_path),
        ]
        
        for path in template_paths_to_try:
            if os.path.exists(path):
                print(f"Loading template: {path}")
                return Document(path)
        
        print(f"Template '{self.template_path}' not found")
        return None
    
    def _flatten_dict(self, d, parent_key='', sep='.'):
        items = {}
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.update(self._flatten_dict(v, new_key, sep=sep))
            else:
                items[new_key] = v
        return items
    
    def _process_template_with_risk_data(self, doc, data):
        """Process template and populate with risk assessment data"""
        
        # Replace simple placeholders
        self._replace_simple_placeholders(doc, data['basic_info'])
        
        # Find and populate the risk assessment table
        self._populate_risk_assessment_table(doc, data['processes'])

        #inventory_table = self._find_work_activity_inventory_table(doc)
        inventory_table = self._find_inventory_table(doc)

        if inventory_table:
            self._fill_work_activity_inventory_table(inventory_table, data['work_activities'])
        else:
            print("Work activity inventory table not found in template")

    
    def _replace_simple_placeholders(self, doc, basic_info):
        """Replace simple text placeholders in the document"""

        # Flatten nested team_data if present
        flat_info = basic_info.copy()
        print("Basic info before flattening:", flat_info)
        if 'team_data' in flat_info and isinstance(flat_info['team_data'], dict):
            flat_info.update(self._flatten_dict(flat_info['team_data'], 'team_data'))
            print("Flattened team_data:", {k: v for k, v in flat_info.items() if k.startswith('team_data.')})
            del flat_info['team_data']  

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in flat_info.items():
                placeholder = f"{{{key}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in flat_info.items():
                        placeholder = f"{{{key}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
    
    def _populate_risk_assessment_table(self, doc, processes_data):
        """Find and populate the risk assessment table"""
        
        # Find the risk assessment table
        risk_table = self._find_risk_assessment_table(doc)
        if not risk_table:
            print("Risk assessment table not found")
            return
        
        print(f"Found risk assessment table with {len(risk_table.rows)} rows and {len(risk_table.columns)} columns")
        
        # Populate the table with data
        self._fill_risk_assessment_table(risk_table, processes_data)
    
    def _find_risk_assessment_table(self, doc):
        """Find the risk assessment table by looking for specific headers"""
        
        for table in doc.tables:
            # Look for identifying text in the table
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text.lower() + " "
            
            print(f"Table {len([t for t in doc.tables if t == table])}: {len(table.columns)} columns")
            print(f"Table text sample: {table_text[:300]}...")  # Show more text
            
            # Check if this looks like the risk assessment table (not the inventory table)
            has_risk_keywords = any(keyword in table_text for keyword in [
                "hazard identification", "risk evaluation", "risk control",
                "possible injury", "ill-health", "existing risk controls", 
                "additional controls", "implementation person", "due date",
                "severity", "likelihood", "rpn", "s l rpn"  # Added this common pattern
            ])
            
            has_required_sections = all(required_keyword in table_text for required_keyword in [
                "hazard identification", "risk evaluation", "risk control"
            ])
            
            is_not_inventory = not any(inventory_keyword in table_text for inventory_keyword in [
                "inventory of work activities", 
                "work activity" + " " + "location",  # More specific pattern
                "location" + " " + "process" + " " + "work activity"  # Specific inventory table pattern
            ])

            print(f"Has risk keywords: {has_risk_keywords}")
            print(f"Has required sections: {has_required_sections}")  
            print(f"Is not inventory: {is_not_inventory}")
            
            # Additional check: look for the specific column structure
            has_correct_structure = ("ref" in table_text and 
                                    "activity" in table_text and 
                                    "hazard" in table_text and
                                    len(table.columns) >= 10)  # Risk assessment tables are wide
            
            if has_risk_keywords and has_required_sections and is_not_inventory and has_correct_structure:
                print("Found risk assessment table!")
                return table
        
        print("Risk assessment table not found")
        return None


    def _apply_table_borders(self, row):
        """Apply borders to table row cells to match existing table formatting"""
        try:
            for cell in row.cells:
                # Get the cell's table cell properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Create borders element
                tcBorders = OxmlElement('w:tcBorders')
                
                # Define border style
                border_attrs = {
                    'w:val': 'single',
                    'w:sz': '4',  # Border width
                    'w:space': '0',
                    'w:color': '000000'  # Black color
                }
                
                # Add all borders (top, left, bottom, right)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border_element = OxmlElement(f'w:{border_name}')
                    for attr_name, attr_value in border_attrs.items():
                        border_element.set(attr_name, attr_value)
                    tcBorders.append(border_element)
                
                tcPr.append(tcBorders)
                
        except Exception as e:
            print(f"Could not apply borders: {e}")

    def _copy_cell_formatting(self, source_cell, target_cell):
        """Copy formatting from source cell to target cell"""
        try:
            # Copy paragraph formatting
            if source_cell.paragraphs and target_cell.paragraphs:
                source_paragraph = source_cell.paragraphs[0]
                target_paragraph = target_cell.paragraphs[0]
                
                # Copy paragraph properties
                if source_paragraph._element.pPr is not None:
                    target_paragraph._element.pPr = source_paragraph._element.pPr
                    
        except Exception as e:
            print(f"Could not copy cell formatting: {e}")

    # Working approach - copy border style from existing row
    def _apply_borders_from_existing_row(self, new_row, existing_row):
        """Copy border formatting from an existing row to a new row"""
        try:
            for i, cell in enumerate(new_row.cells):
                if i < len(existing_row.cells):
                    # Get or create table cell properties for new cell
                    new_tc = cell._tc
                    new_tcPr = new_tc.get_or_add_tcPr()
                    
                    # Get source cell properties
                    source_tc = existing_row.cells[i]._tc
                    source_tcPr = source_tc.tcPr
                    
                    if source_tcPr is not None:
                        # Copy borders if they exist
                        source_borders = source_tcPr.find(qn('w:tcBorders'))
                        if source_borders is not None:
                            # Remove existing borders if any
                            existing_borders = new_tcPr.find(qn('w:tcBorders'))
                            if existing_borders is not None:
                                new_tcPr.remove(existing_borders)
                            
                            # Add copied borders
                            new_tcPr.append(source_borders)
                        
        except Exception as e:
            print(f"Could not copy borders from existing row: {e}")

    # Simplified approach - manually create standard table borders
    def _apply_standard_table_borders(self, row):
        """Apply standard table borders to all cells in a row"""
        try:
            for cell in row.cells:
                # Get or create table cell properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Create borders element
                tcBorders = OxmlElement('w:tcBorders')
                
                # Create each border (top, left, bottom, right)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                
                # Remove existing borders and add new ones
                existing_borders = tcPr.find(qn('w:tcBorders'))
                if existing_borders is not None:
                    tcPr.remove(existing_borders)
                
                tcPr.append(tcBorders)
                        
        except Exception as e:
            print(f"Could not apply standard borders: {e}")

    def _fill_risk_assessment_table(self, table, processes_data):
        """Fill the risk assessment table with actual data - version with merged activity cells for multiple hazards"""
        
        # Find the first empty row (skip header rows)
        data_start_row = self._find_data_start_row(table)
        
        # Clear existing data rows but keep headers
        rows_to_remove = []
        for i in range(data_start_row, len(table.rows)):
            rows_to_remove.append(i)
        
        # Remove existing data rows (in reverse order to avoid index issues)
        for row_idx in reversed(rows_to_remove):
            if row_idx < len(table.rows):
                table._element.remove(table.rows[row_idx]._element)
        
        # Add data for each process
        for process in processes_data:
            # Add process separator row
            process_row = table.add_row()
            self._apply_standard_table_borders(process_row)
            
            if len(process_row.cells) > 0:
                # Merge all cells in the process row
                merged_cell = process_row.cells[0]
                for i in range(1, len(process_row.cells)):
                    merged_cell.merge(process_row.cells[i])
                
                # Set the process name in the merged cell
                merged_cell.text = process['process_name']
                
                # Make process name bold and optionally center it
                for paragraph in merged_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add risk data rows with multi-hazard support and merged activity cells
            for risk_idx, risk in enumerate(process['risks']):
                # Check if this risk has multiple hazards
                hazards = self._extract_hazards_from_risk(risk)
                
                if len(hazards) <= 1:
                    # Single hazard - normal row
                    data_row = table.add_row()
                    self._apply_standard_table_borders(data_row)
                    self._populate_risk_row(data_row, risk)
                    center_columns = [5, 6, 7, 9, 10, 11]
                    self._center_align_columns(data_row, center_columns)
                else:
                    # Multiple hazards - create rows and merge activity cells
                    created_rows = []
                    
                    for hazard_idx, hazard in enumerate(hazards):
                        data_row = table.add_row()
                        self._apply_standard_table_borders(data_row)
                        created_rows.append(data_row)
                        
                        # Create modified risk data for this specific hazard
                        modified_risk = risk.copy()
                        modified_risk['hazard'] = hazard['hazard']
                        modified_risk['possible_injury'] = hazard['possible_injury']
                        
                        # For sub-rows, don't repeat the activity in the activity cell
                        if hazard_idx > 0:
                            modified_risk['activity'] = ''  # Empty for rows that will be merged
                        
                        # Update the reference number for sub-rows
                        if 'ref' in modified_risk:
                            base_ref = modified_risk['ref']
                            modified_risk['ref'] = f"{base_ref}.{hazard_idx + 1}"
                        
                        self._populate_risk_row(data_row, modified_risk)
                        center_columns = [5, 6, 7, 9, 10, 11]
                        self._center_align_columns(data_row, center_columns)
                    
                    # Merge activity cells for all rows of the same activity
                    if len(created_rows) > 1:
                        self._merge_activity_cells(created_rows, risk.get('activity', ''))

    def _merge_activity_cells(self, rows, activity_text):
        """Merge activity cells across multiple rows for the same activity"""
        if not rows or len(rows) < 2:
            return
        
        try:
            # Find the activity column index (usually column 1, but may vary)
            activity_col_idx = self._find_activity_column_index()
            
            # Start with the first row's activity cell
            first_cell = rows[0].cells[activity_col_idx]
            
            # Set the activity text in the first cell
            first_cell.text = activity_text
            
            # Merge subsequent rows' activity cells with the first one
            for i in range(1, len(rows)):
                if activity_col_idx < len(rows[i].cells):
                    # Clear the text in cells that will be merged
                    rows[i].cells[activity_col_idx].text = ''
                    # Merge with the first cell
                    first_cell.merge(rows[i].cells[activity_col_idx])
            
            # Center align the merged activity cell vertically
            self._set_cell_vertical_alignment(first_cell, 'center')
            
        except Exception as e:
            print(f"Error merging activity cells: {e}")

    def _find_activity_column_index(self):
        """Find the index of the activity column. Override this method based on your table structure."""
        # This assumes activity is in column 1 (index 1). 
        # Adjust based on your actual table structure:
        # If your table has: Ref | Activity | Hazard | ... then activity_col_idx = 1
        return 1

    def _set_cell_vertical_alignment(self, cell, alignment='center'):
        """Set vertical alignment for a cell"""
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            
            # Get the cell element
            tc = cell._tc
            
            # Find or create the tcPr element
            tcPr = tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)
            
            # Find or create the vAlign element
            vAlign = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vAlign')
            if vAlign is None:
                vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{alignment}"/>')
                tcPr.append(vAlign)
            else:
                vAlign.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', alignment)
                
        except Exception as e:
            print(f"Error setting vertical alignment: {e}")

    def _extract_hazards_from_risk(self, risk):
        """Extract multiple hazards from a single risk entry"""
        hazards = []
        
        # Check if hazard field contains multiple hazards (separated by common delimiters)
        hazard_text = risk.get('hazard', '')
        injury_text = risk.get('possible_injury', '')
        
        # Try to split by common patterns like "a)", "b)", "1)", "2)", etc.
        if any(pattern in hazard_text.lower() for pattern in ['a)', 'b)', 'c)', 'd)', 'e)']):
            # Split by letter patterns
            hazard_parts = self._split_by_letter_pattern(hazard_text)
            injury_parts = self._split_by_letter_pattern(injury_text)
        elif any(pattern in hazard_text for pattern in ['1)', '2)', '3)', '4)', '5)']):
            # Split by number patterns
            hazard_parts = self._split_by_number_pattern(hazard_text)
            injury_parts = self._split_by_number_pattern(injury_text)
        elif '\n' in hazard_text or ';' in hazard_text:
            # Split by newlines or semicolons
            hazard_parts = [h.strip() for h in hazard_text.replace('\n', ';').split(';') if h.strip()]
            injury_parts = [i.strip() for i in injury_text.replace('\n', ';').split(';') if i.strip()]
        else:
            # Single hazard
            return [{'hazard': hazard_text, 'possible_injury': injury_text}]
        
        # Match hazards with injuries
        for i, hazard in enumerate(hazard_parts):
            injury = injury_parts[i] if i < len(injury_parts) else ''
            hazards.append({
                'hazard': hazard.strip(),
                'possible_injury': injury.strip()
            })
        
        return hazards if hazards else [{'hazard': hazard_text, 'possible_injury': injury_text}]

    def _split_by_letter_pattern(self, text):
        """Split text by letter patterns like 'a)', 'b)', etc."""
        import re
        # Split by pattern like "a)", "b)", etc.
        parts = re.split(r'[a-e]\)\s*', text)
        # Remove empty parts and clean up
        return [part.strip() for part in parts if part.strip()]

    def _split_by_number_pattern(self, text):
        """Split text by number patterns like '1)', '2)', etc."""
        import re
        # Split by pattern like "1)", "2)", etc.
        parts = re.split(r'[1-9]\)\s*', text)
        # Remove empty parts and clean up
        return [part.strip() for part in parts if part.strip()]

    # Alternative simpler approach - copy border style from existing row
    def _apply_borders_from_existing_row(self, new_row, existing_row):
        """Copy border formatting from an existing row to a new row"""
        try:
            for i, cell in enumerate(new_row.cells):
                if i < len(existing_row.cells):
                    # Copy the table cell properties including borders
                    source_tcPr = existing_row.cells[i]._tc.tcPr
                    if source_tcPr is not None:
                        cell._tc.tcPr = source_tcPr
                        
        except Exception as e:
            print(f"Could not copy borders from existing row: {e}")
    
    def _find_data_start_row(self, table):
        """Find where the data rows start (after headers)"""
        # Look for the first row that doesn't contain header text
        for i, row in enumerate(table.rows):
            row_text = ""
            for cell in row.cells:
                row_text += cell.text.lower()
            
            # If row is mostly empty or contains "process", it's likely a data row
            if not any(header in row_text for header in [
                "hazard identification", "risk evaluation", "ref", "activity", "hazard"
            ]) or "process" in row_text:
                return i
        
        # Default to row 2 if we can't determine
        return 2
    
    def _copy_cell_formatting(self, source_cell, target_cell):
        """Copy formatting from source cell to target cell"""
        try:
            # Copy paragraph formatting
            if source_cell.paragraphs and target_cell.paragraphs:
                source_paragraph = source_cell.paragraphs[0]
                target_paragraph = target_cell.paragraphs[0]
                
                # Copy paragraph properties
                if source_paragraph._element.pPr is not None:
                    target_paragraph._element.pPr = source_paragraph._element.pPr
                    
        except Exception as e:
            print(f"Could not copy cell formatting: {e}")
    
    def _populate_risk_row(self, row, risk_data):
        """Populate a single risk row with data"""
        
        # Column mapping based on your table structure
        column_mapping = {
            0: 'ref',
            1: 'activity', 
            2: 'hazard',
            3: 'possible_injury',
            4: 'existing_controls',
            5: 's1',  # Severity
            6: 'l1',  # Likelihood
            7: 'rpn1',  # RPN
            8: 'additional_controls',
            9: 's2',
            10: 'l2', 
            11: 'rpn2',
            12: 'implementation_person',
            13: 'due_date',
            14: 'remarks'
        }
        
        # Fill the row with data
        for col_idx, field_name in column_mapping.items():
            if col_idx < len(row.cells) and field_name in risk_data:
                row.cells[col_idx].text = str(risk_data[field_name])
    
    def _merge_cells_for_process_name(self, row, total_columns):
        """Merge cells for process name row"""
        try:
            # This is a simplified version - actual cell merging in python-docx
            # requires more complex XML manipulation
            if len(row.cells) > 1:
                # Clear other cells in the row
                for i in range(1, min(len(row.cells), total_columns)):
                    row.cells[i].text = ""
        except Exception as e:
            print(f"Could not merge cells: {e}")
    
    def _add_dynamic_columns_if_needed(self, table, risk_data):
        """Add additional columns if the data requires them"""
        
        if not risk_data:
            return
        
        # Check if we need more columns
        sample_risk = risk_data[0]['risks'][0] if risk_data[0]['risks'] else {}
        required_columns = len(sample_risk.keys())
        current_columns = len(table.columns)
        
        if required_columns > current_columns:
            # Add additional columns
            for _ in range(required_columns - current_columns):
                table.add_column(Inches(0.8))
            
            print(f"Added {required_columns - current_columns} additional columns")

    def _apply_table_borders(self, row):
        """Apply borders to table row cells to match existing table formatting"""
        try:
            for cell in row.cells:
                # Get the cell's table cell properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Create borders element
                tcBorders = OxmlElement('w:tcBorders')
                
                # Define border style
                border_attrs = {
                    'w:val': 'single',
                    'w:sz': '4',  # Border width
                    'w:space': '0',
                    'w:color': '000000'  # Black color
                }
                
                # Add all borders (top, left, bottom, right)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border_element = OxmlElement(f'w:{border_name}')
                    for attr_name, attr_value in border_attrs.items():
                        border_element.set(attr_name, attr_value)
                    tcBorders.append(border_element)
                
                tcPr.append(tcBorders)
                
        except Exception as e:
            print(f"Could not apply borders: {e}")

   
    def _find_inventory_table(self, doc):
        """Find the inventory of work activities DATA table (not the header table)"""
        
        for i, table in enumerate(doc.tables):
            # Extract all text from the table
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text.lower().strip() + " "
            
            print(f"Checking table {i+1}: {len(table.columns)} columns, {len(table.rows)} rows")
            print(f"Table text sample: {table_text[:200]}...")
            
            # Look for the DATA table specifically (not the header table)
            # The data table should have the column headers: Ref, Location, Process, Work Activity, Remarks
            has_ref_column = "ref" in table_text
            has_location_column = "location" in table_text
            has_process_column = "process" in table_text
            has_work_activity_column = "work activity" in table_text
            has_remarks_column = "remarks" in table_text
            
            # Check if this table has numbered rows (1, 2, 3, 4, etc.)
            has_numbered_rows = False
            for row in table.rows:
                first_cell = row.cells[0].text.strip()
                if first_cell.isdigit():
                    has_numbered_rows = True
                    break
            
            # Check table structure - should be exactly 5 columns for data table
            has_correct_width = len(table.columns) == 5
            
            # This should NOT be the header table (which has "Reference Number", "Division", etc.)
            is_not_header_table = not any(keyword in table_text for keyword in [
                "reference number", "division", "{title}", "{division}",
                "prs ra repository", "next running number"
            ])
            
            # Ensure it's NOT a risk assessment table
            is_not_risk_table = not any(keyword in table_text for keyword in [
                "hazard identification", "risk evaluation", "severity", "likelihood", 
                "rpn", "existing risk controls", "additional controls", "implementation person"
            ])
            
            print(f"  Has data columns: ref={has_ref_column}, location={has_location_column}, process={has_process_column}")
            print(f"  Has work_activity={has_work_activity_column}, remarks={has_remarks_column}")
            print(f"  Has numbered rows: {has_numbered_rows}")
            print(f"  Correct width (5 cols): {has_correct_width}")
            print(f"  Not header table: {is_not_header_table}")
            print(f"  Not risk table: {is_not_risk_table}")
            
            # Match criteria - must have all data columns and correct structure
            if (has_ref_column and has_location_column and has_process_column and 
                has_work_activity_column and has_remarks_column and 
                has_correct_width and is_not_header_table and is_not_risk_table):
                print(f"Found inventory DATA table at index {i+1}")
                return table
        
        print("Inventory data table not found")
        return None

    def _populate_work_activity_inventory_table(self, doc, activities_data):
        """Find and populate the work activity inventory table"""
        
        inventory_table = self._find_inventory_table(doc)
        
        if not inventory_table:
            print("Work inventory table not found")
            return False
        
        print(f"Found inventory table with {len(inventory_table.rows)} rows and {len(inventory_table.columns)} columns")
        
        # Fill the table with data
        self._fill_work_activity_inventory_table(inventory_table, activities_data)
        return True

    def _apply_clean_inventory_formatting(self, row):
        """Apply clean, consistent formatting to inventory rows"""
        
        try:
            # Apply borders first
            self._apply_standard_table_borders(row)
            
            # Set consistent font and alignment for each cell
            for i, cell in enumerate(row.cells):
                # Clear any existing content
                cell.text = ""
                
                # Ensure we have at least one paragraph
                if not cell.paragraphs:
                    cell.add_paragraph()
                
                # Set paragraph alignment
                for paragraph in cell.paragraphs:
                    if i == 0:  # First column (Ref) - center aligned
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:  # Other columns - left aligned
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Set font properties
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
            
            print(f"Applied clean formatting to row")
            
        except Exception as e:
            print(f"Warning: Could not apply formatting: {e}")


    def _find_inventory_data_start_row(self, table):
        """Find where the actual data rows start in the inventory table with enhanced debugging"""
        
        print(f"DEBUG: Table has {len(table.rows)} rows, {len(table.columns)} columns")
        
        # Print all rows to understand the structure
        for i, row in enumerate(table.rows):
            row_text = ""
            for j, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_text += f"[{j}:'{cell_text}'] "
            print(f"Row {i}: {row_text}")
        
        # Look for the first numbered row (1, 2, 3, etc.)
        for i, row in enumerate(table.rows):
            first_cell_text = row.cells[0].text.strip()
            print(f"Checking row {i}, first cell: '{first_cell_text}'")
            
            # Skip header row and look for numbered data rows
            if first_cell_text == "1":
                print(f"Found data start at row {i} (first numbered row)")
                return i
            elif first_cell_text.isdigit():
                print(f"Found data start at row {i} (numbered row: {first_cell_text})")
                return i
        
        # Fallback - look for the row after the header row with column names
        for i, row in enumerate(table.rows):
            row_text = ""
            for cell in row.cells:
                row_text += cell.text.lower().strip() + " "
            
            print(f"Checking row {i} for headers: '{row_text}'")
            
            # If this is the header row, data starts next
            if ("ref" in row_text and "location" in row_text and 
                "process" in row_text and "work activity" in row_text):
                print(f"Found header row at {i}, data starts at {i+1}")
                return i + 1
        
        # Last resort - assume row 1 is data (after potential header at row 0)
        print("Using fallback: data starts at row 1")
        return 1

    def _populate_inventory_row_clean(self, row, activity_data, ref_number):
        """Populate inventory row with clean approach"""
        
        try:
            # Data to populate
            cell_data = [
                str(ref_number),
                activity_data.get('location', ''),
                activity_data.get('process', ''),
                activity_data.get('work_activity', ''),
                activity_data.get('remarks', '')
            ]
            
            # Populate each cell
            for i, data in enumerate(cell_data):
                if i < len(row.cells):
                    cell = row.cells[i]
                    
                    # Clear existing content
                    cell.text = ""
                    
                    # Add content
                    if data:
                        cell.text = data
                    
                    # Set alignment and formatting
                    if cell.paragraphs:
                        paragraph = cell.paragraphs[0]
                        
                        # Set alignment
                        if i == 0:  # Ref column - center
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:  # Other columns - left
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Set font
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
            
            print(f"Populated row {ref_number}: {activity_data.get('location', '')} - {activity_data.get('work_activity', '')}")
            
        except Exception as e:
            print(f"ERROR populating row {ref_number}: {e}")
            # Fallback - just set text
            for i, data in enumerate(cell_data):
                if i < len(row.cells):
                    row.cells[i].text = data

    def _copy_inventory_row_formatting(self, source_row, target_row):
        """Copy formatting from source row to target row"""
        
        try:
            # Ensure target row has enough cells
            while len(target_row.cells) < len(source_row.cells):
                target_row.add_cell()
            
            # Copy formatting for each cell
            for i, (source_cell, target_cell) in enumerate(zip(source_row.cells, target_row.cells)):
                # Copy paragraph formatting
                if source_cell.paragraphs and target_cell.paragraphs:
                    source_para = source_cell.paragraphs[0]
                    target_para = target_cell.paragraphs[0]
                    
                    # Copy alignment
                    if source_para.alignment is not None:
                        target_para.alignment = source_para.alignment
                    
                    # Copy font formatting from runs
                    if source_para.runs:
                        # Clear existing runs in target
                        for run in target_para.runs:
                            run.clear()
                        
                        # Copy from source
                        for source_run in source_para.runs:
                            target_run = target_para.add_run("")
                            if source_run.font.name:
                                target_run.font.name = source_run.font.name
                            if source_run.font.size:
                                target_run.font.size = source_run.font.size
                            if source_run.font.bold is not None:
                                target_run.font.bold = source_run.font.bold
                            if source_run.font.italic is not None:
                                target_run.font.italic = source_run.font.italic
                
                # Copy cell properties (borders, shading)
                self._copy_cell_properties(source_cell, target_cell)
            
            # Apply table borders
            self._apply_standard_table_borders(target_row)
            
        except Exception as e:
            print(f"Warning: Could not copy all formatting: {e}")
            # Fallback to basic formatting
            self._apply_inventory_row_formatting(target_row)

    def _copy_cell_properties(self, source_cell, target_cell):
        """Copy cell-level properties like borders and shading"""
        
        try:
            if hasattr(source_cell, '_element') and hasattr(target_cell, '_element'):
                source_tcPr = source_cell._element.get_or_add_tcPr()
                target_tcPr = target_cell._element.get_or_add_tcPr()
                
                # Copy shading
                source_shd = source_tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if source_shd is not None:
                    target_shd = OxmlElement('w:shd')
                    target_shd.set(qn('w:fill'), source_shd.get(qn('w:fill'), 'auto'))
                    target_shd.set(qn('w:val'), source_shd.get(qn('w:val'), 'clear'))
                    target_tcPr.append(target_shd)
                
                # Copy borders
                source_tcBorders = source_tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
                if source_tcBorders is not None:
                    target_tcBorders = OxmlElement('w:tcBorders')
                    target_tcBorders[:] = source_tcBorders[:]
                    target_tcPr.append(target_tcBorders)
        
        except Exception as e:
            print(f"Warning: Could not copy cell properties: {e}")

    def _apply_inventory_row_formatting(self, row):
        """Apply standard formatting to inventory row"""
        
        try:
            for i, cell in enumerate(row.cells):
                # Apply font formatting
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                    
                    # Center align the first column (Ref)
                    if i == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Apply borders
            self._apply_standard_table_borders(row)
        
        except Exception as e:
            print(f"Warning: Could not apply formatting: {e}")

    def _apply_standard_table_borders(self, row):
        """Apply standard table borders to a row"""
        
        try:
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                
                # Create border element
                tcBorders = OxmlElement('w:tcBorders')
                
                # Define border style
                border_attrs = {
                    'w:val': 'single',
                    'w:sz': '4',
                    'w:space': '0',
                    'w:color': '000000'
                }
                
                # Add all borders
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border_el = OxmlElement(f'w:{border_name}')
                    for attr, value in border_attrs.items():
                        border_el.set(qn(attr), value)
                    tcBorders.append(border_el)
                
                tcPr.append(tcBorders)
        
        except Exception as e:
            print(f"Warning: Could not apply borders: {e}")


    
    def _fill_work_activity_inventory_table(self, table, activities_data):
        """Fill the work activity inventory table with merged cells for same processes"""
        
        print(f"Table structure: {len(table.rows)} rows, {len(table.columns)} columns")
        
        header_row_index = 0
        data_start_row = 1
        
        # Clear all existing data rows (keep only header)
        rows_to_remove = []
        for i in range(data_start_row, len(table.rows)):
            rows_to_remove.append(i)
        
        print(f"Removing {len(rows_to_remove)} existing template rows...")
        for row_idx in reversed(rows_to_remove):
            if row_idx < len(table.rows):
                table._element.remove(table.rows[row_idx]._element)
        
        print(f"After cleanup: {len(table.rows)} rows remain")
        
        # Group activities by process and location
        grouped_data = self._group_activities_by_process(activities_data)
        
        # Add rows and populate with merging
        self._populate_grouped_inventory_table(table, grouped_data)
        
        print("Table population with merging complete")

    def _group_activities_by_process(self, activities):
        """Group activities by process and location, handling None values"""
        try:
            print(f"=== DEBUG: group_activities_by_process ===")
            print(f"Input activities: {activities}")
            
            if not activities:
                print("No activities provided")
                return {}
            
            # Ensure activities is a list
            if not isinstance(activities, list):
                print(f"Activities is not a list: {type(activities)}")
                return {}
            
            # Clean and validate the data first
            cleaned_activities = []
            for activity in activities:
                if not isinstance(activity, dict):
                    print(f"Skipping invalid activity (not dict): {activity}")
                    continue
                
                # Handle None values by providing defaults
                cleaned_activity = {
                    'location': activity.get('location', 'Unknown Location'),
                    'process': activity.get('process', 'Unknown Process'),
                    'hazard': activity.get('hazard', 'Unknown Hazard'),
                    'existing_control': activity.get('existing_control', 'None specified'),
                    'likelihood': activity.get('likelihood', 'Unknown'),
                    'consequence': activity.get('consequence', 'Unknown'),
                    'risk_rating': activity.get('risk_rating', 'Unknown'),
                    'additional_control': activity.get('additional_control', 'None specified'),
                    'residual_likelihood': activity.get('residual_likelihood', 'Unknown'),
                    'residual_consequence': activity.get('residual_consequence', 'Unknown'),
                    'residual_risk_rating': activity.get('residual_risk_rating', 'Unknown'),
                    'person_responsible': activity.get('person_responsible', 'TBD'),
                    'target_date': activity.get('target_date', 'TBD'),
                    'work_activity': activity.get('work_activity', activity.get('hazard', 'Unknown Activity')),
                    'remarks': activity.get('remarks', activity.get('existing_control', 'None specified'))
           
                }
                cleaned_activities.append(cleaned_activity)
            
            print(f"Cleaned activities: {cleaned_activities}")
            
            # Sort the cleaned activities (now all values are strings, not None)
            cleaned_activities.sort(key=lambda x: (x['location'], x['process']))
            
            # Group by location and process
            result = {}
            for activity in cleaned_activities:
                location = activity['location']
                process = activity['process']
                
                if location not in result:
                    result[location] = {}
                if process not in result[location]:
                    result[location][process] = []
                
                result[location][process].append(activity)
            
            print(f"Grouped result: {result}")
            return result
            
        except Exception as e:
            print(f"Error in group_activities_by_process: {e}")
            import traceback
            traceback.print_exc()
            return {}

    def _populate_grouped_inventory_table(self, table, grouped_data):
        """Populate table with grouped data and merge cells"""
        
        try:
            print(f"=== DEBUG: _populate_grouped_inventory_table ===")
            print(f"grouped_data type: {type(grouped_data)}")
            print(f"grouped_data content: {grouped_data}")
            
            # Convert nested dictionary to expected list format if needed
            if isinstance(grouped_data, dict):
                print("Converting nested dict to list format...")
                list_data = []
                for location_name, processes in grouped_data.items():
                    if isinstance(processes, dict):
                        for process_name, activities in processes.items():
                            if isinstance(activities, list):
                                # Ensure activities have the required structure
                                processed_activities = []
                                for activity in activities:
                                    if isinstance(activity, dict):
                                        processed_activity = {
                                            'work_activity': activity.get('hazard', ''),
                                            'remarks': activity.get('existing_control', ''),
                                            # Keep original fields too
                                            **activity
                                        }
                                        processed_activities.append(processed_activity)
                                
                                group_data = {
                                    'location': location_name,
                                    'process': process_name,
                                    'activities': processed_activities
                                }
                                list_data.append(group_data)
                grouped_data = list_data
            
            if not grouped_data:
                print("No grouped data to populate")
                return
                
            ref_counter = 1
            
            for group in grouped_data:
                print(f"Processing group: {group}")
                
                # Handle the case where group might be a string (the original error)
                if isinstance(group, str):
                    print(f"ERROR: Expected dict for group, got string: {group}")
                    continue
                
                if not isinstance(group, dict):
                    print(f"ERROR: Expected dict for group, got {type(group)}")
                    continue
                
                location = group.get('location', 'Unknown Location')
                process = group.get('process', 'Unknown Process')
                activities = group.get('activities', [])
                
                print(f"Group - Location: {location}, Process: {process}, Activities: {len(activities)}")
                
                # Track the starting row for this group
                group_start_row = len(table.rows)
                
                # Add rows for each activity in this group
                for i, activity in enumerate(activities):
                    new_row = table.add_row()
                    self._apply_clean_inventory_formatting(new_row)

                    ref_display = str(ref_counter) if i == 0 else ""
                    
                    # Populate the row
                    self._populate_merged_inventory_row(
                        new_row, 
                        ref_display, 
                        location if i == 0 else "",  # Only show location in first row
                        process if i == 0 else "",   # Only show process in first row
                        activity.get('work_activity', ''),
                        activity.get('remarks', '')
                    )
                    
                ref_counter += 1
                
                # Merge cells for location and process columns if multiple activities
                if len(activities) > 1:
                    self._merge_process_cells(table, group_start_row, len(activities))
            
            print(f"Populated {ref_counter - 1} total groups")
            
        except Exception as e:
            print(f"Error in _populate_grouped_inventory_table: {e}")
            import traceback
            traceback.print_exc()

    def _populate_merged_inventory_row(self, row, ref_display, location, process, work_activity, remarks, is_single_activity=False):
        """Populate a single inventory row (modified for merged layout)"""
        
        try:
            print(f"Populating row - ref: {ref_display}, location: {location}, process: {process}")
            
            # Handle None values
            ref_display = str(ref_display) if ref_display else ""
            location = str(location) if location else ""
            process = str(process) if process else ""
            work_activity = str(work_activity) if work_activity else ""
            remarks = str(remarks) if remarks else ""
            
            # Data to populate (location and process might be empty for merged rows)
            cell_data = [
                ref_display,
                location,
                process,
                work_activity,
                remarks
            ]
            
            # Populate each cell
            for i, data in enumerate(cell_data):
                if i < len(row.cells):
                    cell = row.cells[i]
                    
                    # Clear existing content
                    cell.text = ""
                    
                    # Add content
                    if data:
                        cell.text = data
                    
                    # Set alignment and formatting
                    if cell.paragraphs:
                        paragraph = cell.paragraphs[0]
                        
                        # Set alignment based on column and content
                        if i == 0:  # Ref column - always center
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif i == 1:  # Location column
                            if is_single_activity:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif i == 2:  # Process column - always left
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:  # Work Activity and Remarks columns - left
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Set font
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)

            process_display = process if process else "[continued]"
            print(f"Successfully populated row {ref_display}: {location} - {process_display} - {work_activity}")
                    
        except Exception as e:
            print(f"ERROR populating row {ref_display}: {e}")
            import traceback
            traceback.print_exc()

    def _merge_process_cells(self, table, start_row_index, num_activities):
        """Merge location and process cells for activities belonging to same process"""
        
        try:

            # Merge ref column (column 0) - same reference number for all activities in process
            if num_activities > 1:
                ref_cell = table.rows[start_row_index].cells[0]
                for i in range(1, num_activities):
                    merge_cell = table.rows[start_row_index + i].cells[0]
                    ref_cell.merge(merge_cell)

            # Merge location column (column 1)
            if num_activities > 1:
                location_cell = table.rows[start_row_index].cells[1]
                for i in range(1, num_activities):
                    merge_cell = table.rows[start_row_index + i].cells[1]
                    location_cell.merge(merge_cell)
            
            # Merge process column (column 2)
            if num_activities > 1:
                process_cell = table.rows[start_row_index].cells[2]
                for i in range(1, num_activities):
                    merge_cell = table.rows[start_row_index + i].cells[2]
                    process_cell.merge(merge_cell)
            
            # Set vertical alignment for merged cells
            self._set_merged_cell_alignment(table.rows[start_row_index].cells[0], 'center')  # Ref - center
            self._set_merged_cell_alignment(table.rows[start_row_index].cells[1])
            self._set_merged_cell_alignment(table.rows[start_row_index].cells[2])
            
            print(f"Merged cells for {num_activities} activities starting at row {start_row_index}")
            
        except Exception as e:
            print(f"ERROR merging cells: {e}")

    def _set_merged_cell_alignment(self, cell, horizontal_align='center'):
        """Set vertical alignment for merged cells"""
        
        try:
            # Set vertical alignment to center
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            # Create vertical alignment element
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            
            # Set paragraph alignment
            for paragraph in cell.paragraphs:
                if horizontal_align == 'center':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif horizontal_align == 'left':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif horizontal_align == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        except Exception as e:
            print(f"Warning: Could not set merged cell alignment: {e}")


    def _replace_page_placeholders_with_fields(self, doc):
        """Replace {PAGE} and {NUMPAGES} with actual Word field codes"""
        for paragraph in doc.paragraphs:
            if '{PAGE}' in paragraph.text or '{NUMPAGES}' in paragraph.text:
                self._replace_fields_in_paragraph(paragraph)

