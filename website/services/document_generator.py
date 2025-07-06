
import os
import shutil
from docx import Document
from docx.shared import Inches
from datetime import datetime
from jinja2 import Template
import xml.etree.ElementTree as ET
from models import *

# Alternative approach using python-docx-template (recommended)
class DocxTemplateGenerator:
    """
    Alternative implementation using python-docx-template library
    This is easier and more robust for complex templates
    """
    
    def __init__(self, template_path="Risk_Assessment_Form_Template.docx"):
        self.template_path = template_path
        
    def generate_word_document(self, assessment_id: str, output_path: str = None):
        """Generate document using docx-template"""
        try:
            # You'll need to install: pip install python-docx-template
            from docxtpl import DocxTemplate
            
            if output_path is None:
                output_path = f"risk_assessment_{assessment_id}.docx"
            
            # Load template
            doc = DocxTemplate(self.template_path)
            
            # Get data
            header, rows = self.get_assessment(assessment_id)
            
            # Prepare context
            context = self.prepare_template_context(header, rows)
            
            # Render template
            doc.render(context)
            
            # Save
            doc.save(output_path)
            
            print(f"Document generated successfully: {output_path}")
            return output_path
            
        except ImportError:
            print("python-docx-template not installed. Please install it with:")
            print("pip install python-docx-template")
            return None
        except Exception as e:
            print(f"Error generating document: {e}")
            return None
    
    def prepare_template_context(self, header, rows):
        """Prepare context for docx-template"""
        def safe_str(value):
            return str(value) if value is not None else ""
        
        # Header context
        header_context = {}
        if header:
            header_context = {
                'reference_number': safe_str(header.reference_number),
                'division': safe_str(header.division),
                'title': safe_str(header.title),
                'location': safe_str(header.location),
                'ra_leader': safe_str(header.ra_leader),
                'ra_team': safe_str(header.ra_team),
                'approved_by': safe_str(header.approved_by),
                'signature': safe_str(header.signature),
                'designation': safe_str(header.designation),
                'date': safe_str(header.date),
                'last_review_date': safe_str(header.last_review_date),
                'next_review_date': safe_str(header.next_review_date)
            }
        
        # Process rows
        processes = []
        if rows:
            current_process = None
            process_rows = []
            
            for row in rows:
                process_name = safe_str(row.process_name) if hasattr(row, 'process_name') and row.process_name else "General Process"
                
                if current_process != process_name:
                    if current_process is not None:
                        processes.append({
                            'name': current_process,
                            'rows': process_rows
                        })
                    current_process = process_name
                    process_rows = []
                
                process_rows.append({
                    'ref': safe_str(row.ref),
                    'activity': safe_str(row.activity),
                    'hazard': safe_str(row.hazard),
                    'possible_injury': safe_str(row.possible_injury),
                    'existing_controls': safe_str(row.existing_controls),
                    'severity_initial': safe_str(row.severity_initial),
                    'likelihood_initial': safe_str(row.likelihood_initial),
                    'rpn_initial': safe_str(row.rpn_initial),
                    'additional_controls': safe_str(row.additional_controls),
                    'severity_final': safe_str(row.severity_final),
                    'likelihood_final': safe_str(row.likelihood_final),
                    'rpn_final': safe_str(row.rpn_final),
                    'implementation_person': safe_str(row.implementation_person),
                    'due_date': safe_str(row.due_date),
                    'remarks': safe_str(row.remarks)
                })
            
            if current_process is not None:
                processes.append({
                    'name': current_process,
                    'rows': process_rows
                })
        
        return {
            'header': header_context,
            'processes': processes,
            'generation_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
    
    def get_assessment(self, assessment_id):
        """Actual implementation for your data retrieval"""
        # Replace with your actual database query
        header = Form.query.get(assessment_id)
        rows = header.assessment_rows if header else []
        return header, rows
    
    def generate_with_python_docx(self, assessment_id: str, output_path: str = None):
        """Fallback method using python-docx only - Testing with hardcoded values"""
        try:
            if output_path is None:
                output_path = f"risk_assessment_{assessment_id}.docx"
            
            # Create new document
            doc = Document()
            
            # Add title
            doc.add_heading('Risk Assessment Form', 0)
            
            # Add division section with hardcoded test values
            doc.add_heading('Division Information', 1)
            
            # Create division table
            division_table = doc.add_table(rows=0, cols=2)
            division_table.style = 'Table Grid'
            
            # Hardcoded test data for division section
            division_data = [
                ('Division', 'Engineering & Operations'),
                ('Department', 'Safety & Quality Assurance'),
                ('Unit', 'Risk Management Team'),
                ('Location', 'Main Plant - Building A'),
                ('Supervisor', 'John Smith'),
                ('Date Created', '2024-01-15'),
                ('Status', 'Active')
            ]
            
            # Add division data to table
            for label, value in division_data:
                row = division_table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(value)
                
                # Make first column bold
                row.cells[0].paragraphs[0].runs[0].bold = True
            
            # Add some spacing
            doc.add_paragraph()
            
            # Add a simple note
            doc.add_paragraph("Note: This is a test document with hardcoded values for the division section.")
            
            # Save document
            doc.save(output_path)
            print(f"Test document generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"Error generating test document: {e}")
            return None

# Usage example
if __name__ == "__main__":
    
    # Method 2: Using python-docx-template (recommended)
    generator2 = DocxTemplateGenerator("Risk_Assessment_Form_Template.docx")
    
    print("Template generators ready!")
    print("\nRecommended approach:")
    print("1. Install python-docx-template: pip install python-docx-template")
    print("2. Edit your Word template to include Jinja2 variables like {{ header.reference_number }}")
    print("3. Use DocxTemplateGenerator for easier template management")